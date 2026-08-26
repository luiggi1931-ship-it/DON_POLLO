import os
import tempfile
import subprocess
from datetime import datetime
from flask import Blueprint, jsonify, send_file, request, current_app, flash, redirect, url_for
from flask_login import login_required
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from models.lote import Lote
from models.lectura_sensor import LecturaSensor
from models.mortalidad import Mortalidad
from models.vacunacion import Vacunacion
from models.vacuna import Vacuna

exportacion_bp = Blueprint('exportacion', __name__)

@exportacion_bp.route('/api/reportes/exportar/<formato>/<int:id_lote>')
@login_required
def exportar_reporte(formato, id_lote):
    lote = Lote.query.get_or_404(id_lote)
    
    try:
        if formato == 'excel':
            return exportar_excel(lote)
        elif formato == 'word':
            return exportar_word(lote)
        elif formato == 'pdf':
            return exportar_pdf_fpdf(lote)
        else:
            return jsonify({'error': 'Formato no soportado'}), 400
    except Exception as e:
        import traceback
        traceback.print_exc()
        flash(f'Error al generar el reporte {formato}: {str(e)}', 'error')
        return redirect(url_for('lotes.lotes'))


def exportar_excel(lote):
    # Recopilar telemetría
    lecturas = LecturaSensor.query.filter_by(id_lote=lote.id_lote_).order_by(LecturaSensor.fecha_hora.asc()).all()
    data_telemetria = []
    for lec in lecturas:
        data_telemetria.append({
            'Fecha y Hora': lec.fecha_hora.strftime('%Y-%m-%d %H:%M:%S'),
            'Temperatura (°C)': lec.temperatura,
            'Humedad (%)': lec.humedad,
            'Amoníaco (ppm)': lec.amoniaco,
            'Iluminación (lux)': lec.iluminacion,
            'Agua (L)': lec.agua,
            'Alimento (kg)': lec.alimento
        })
    df_telemetria = pd.DataFrame(data_telemetria)
    
    # Recopilar mortalidad
    bajas = Mortalidad.query.filter_by(id_lote_=lote.id_lote_).all()
    data_mortalidad = []
    for b in bajas:
        data_mortalidad.append({
            'Fecha': b.fecha.strftime('%Y-%m-%d'),
            'Causa': b.causa,
            'Cantidad': b.cantidad
        })
    df_mortalidad = pd.DataFrame(data_mortalidad)

    # Recopilar vacunación
    from sqlalchemy.orm import joinedload
    vacunas = Vacunacion.query.filter_by(id_lote=lote.id_lote_).options(joinedload(Vacunacion.vacuna)).all()
    data_vacunas = []
    for v in vacunas:
        data_vacunas.append({
            'Fecha': v.fecha_aplicacion.strftime('%Y-%m-%d') if v.fecha_aplicacion else '',
            'Vacuna': v.vacuna.nombre if v.vacuna else f'ID: {v.id_vacuna}',
            'Enfermedad Objetivo': v.vacuna.enfermedad_objetivo if v.vacuna else '',
            'Estado': v.estado,
            'Observaciones': v.observaciones or ''
        })
    df_vacunas = pd.DataFrame(data_vacunas)

    # Crear Excel
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f'Reporte_Lote_{lote.id_lote_}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
    
    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        if not df_telemetria.empty:
            df_telemetria.to_excel(writer, sheet_name='Telemetría', index=False)
        else:
            pd.DataFrame({'Mensaje': ['No hay datos de telemetría']}).to_excel(writer, sheet_name='Telemetría', index=False)
            
        if not df_mortalidad.empty:
            df_mortalidad.to_excel(writer, sheet_name='Mortalidad', index=False)
        else:
            pd.DataFrame({'Mensaje': ['No hay registros de mortalidad']}).to_excel(writer, sheet_name='Mortalidad', index=False)

        if not df_vacunas.empty:
            df_vacunas.to_excel(writer, sheet_name='Vacunación', index=False)
        else:
            pd.DataFrame({'Mensaje': ['No hay registros de vacunación']}).to_excel(writer, sheet_name='Vacunación', index=False)

    return send_file(file_path, as_attachment=True)


def exportar_word(lote):
    doc = Document()
    
    # Título
    title = doc.add_heading(f'Informe de Gestión - Lote #{lote.id_lote_}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha de emisión: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    estado = "Activo" if lote.estado_activo_cerrado == 1 else "Cerrado"
    p_resumen = doc.add_paragraph(f"El lote #{lote.id_lote_} ingresó a la granja el {lote.fecha_ingreso.strftime('%Y-%m-%d')} proveniente del proveedor {lote.proveedor}. ")
    p_resumen.add_run(f"Actualmente se encuentra en estado {estado} con una edad de {lote.edad_dias} días. ")
    p_resumen.add_run(f"El lote cuenta con un saldo de {lote.saldo_actual} aves (de las {lote.cantidad_inicial} originales), representando una mortalidad del {lote.mortalidad_porcentaje}%.")
    
    # Datos Generales
    doc.add_heading('2. Datos Generales del Lote', level=1)
    table_datos = doc.add_table(rows=4, cols=2)
    table_datos.style = 'Light Shading Accent 1'
    datos = [
        ("Tipo de Ave:", lote.tipo_ave),
        ("Proveedor:", lote.proveedor),
        ("Peso Inicial:", f"{lote.peso_inicial} g"),
        ("Edad:", f"{lote.edad_dias} días")
    ]
    for i, (k, v) in enumerate(datos):
        table_datos.cell(i, 0).text = k
        table_datos.cell(i, 1).text = str(v)
    doc.add_paragraph() # Spacing
        
    # Telemetría Promedios
    doc.add_heading('3. Resumen Ambiental (Promedios Históricos)', level=1)
    lecturas = LecturaSensor.query.filter_by(id_lote=lote.id_lote_).all()
    if lecturas:
        avg_temp = round(sum(l.temperatura for l in lecturas) / len(lecturas), 1)
        avg_hum = round(sum(l.humedad for l in lecturas) / len(lecturas), 1)
        avg_amon = round(sum(l.amoniaco for l in lecturas) / len(lecturas), 1)
        avg_luz = round(sum(l.iluminacion for l in lecturas) / len(lecturas), 1)
        tot_agua = round(sum(l.agua for l in lecturas), 1)
        tot_ali = round(sum(l.alimento for l in lecturas), 1)
        
        table_env = doc.add_table(rows=6, cols=2)
        table_env.style = 'Light List Accent 1'
        env_data = [
            ("Temperatura Promedio", f"{avg_temp} °C"),
            ("Humedad Promedio", f"{avg_hum} %"),
            ("Amoníaco Promedio", f"{avg_amon} ppm"),
            ("Iluminación Promedio", f"{avg_luz} lux"),
            ("Consumo Total de Agua Registrado", f"{tot_agua} L"),
            ("Consumo Total de Alimento Registrado", f"{tot_ali} kg")
        ]
        for i, (k, v) in enumerate(env_data):
            table_env.cell(i, 0).text = k
            table_env.cell(i, 1).text = str(v)
    else:
        doc.add_paragraph("No hay datos de telemetría suficientes para este lote.")
    doc.add_paragraph()
    
    # Vacunaciones
    doc.add_heading('4. Historial de Vacunación', level=1)
    vacunas = Vacunacion.query.filter_by(id_lote=lote.id_lote_).all()
    if vacunas:
        table_vac = doc.add_table(rows=1, cols=4)
        table_vac.style = 'Medium Shading 1 Accent 1'
        hdr_cells = table_vac.rows[0].cells
        hdr_cells[0].text = 'Fecha'
        hdr_cells[1].text = 'Vacuna'
        hdr_cells[2].text = 'Objetivo'
        hdr_cells[3].text = 'Estado'
        for v in vacunas:
            row_cells = table_vac.add_row().cells
            row_cells[0].text = v.fecha_aplicacion.strftime('%Y-%m-%d') if v.fecha_aplicacion else 'N/A'
            row_cells[1].text = v.vacuna.nombre if v.vacuna else 'Desconocida'
            row_cells[2].text = v.vacuna.enfermedad_objetivo if v.vacuna else 'Desconocido'
            row_cells[3].text = v.estado
    else:
        doc.add_paragraph("No se han registrado aplicaciones de vacunas para este lote.")
    doc.add_paragraph()
    
    # Mortalidad
    doc.add_heading('5. Detalle de Mortalidad', level=1)
    bajas = Mortalidad.query.filter_by(id_lote_=lote.id_lote_).all()
    if bajas:
        table_mor = doc.add_table(rows=1, cols=3)
        table_mor.style = 'Medium Shading 1 Accent 1'
        hdr_cells = table_mor.rows[0].cells
        hdr_cells[0].text = 'Fecha'
        hdr_cells[1].text = 'Causa'
        hdr_cells[2].text = 'Cantidad'
        for b in bajas:
            row_cells = table_mor.add_row().cells
            row_cells[0].text = b.fecha.strftime('%Y-%m-%d')
            row_cells[1].text = b.causa
            row_cells[2].text = str(b.cantidad)
    else:
        doc.add_paragraph("No se han registrado bajas en este lote.")
    
    # Espacios para notas gerenciales
    doc.add_page_break()
    doc.add_heading('6. Observaciones y Notas', level=1)
    for _ in range(5):
        p = doc.add_paragraph('____________________________________________________________________')
    
    doc.add_heading('Firma del Administrador / Veterinario', level=1)
    doc.add_paragraph('\n\n\n_________________________________')
    doc.add_paragraph('Nombre y Firma')

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f'Informe_Gestion_Lote_{lote.id_lote_}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(file_path)
    
    return send_file(file_path, as_attachment=True)


def exportar_pdf_fpdf(lote):
    pdf = FPDF()
    pdf.add_page()
    
    # Colores corporativos
    COLOR_PRIMARIO = (29, 53, 87) # Azul oscuro
    COLOR_SECUNDARIO = (69, 123, 157) # Azul medio
    COLOR_GRIS = (241, 245, 249) # Gris claro fondo tabla
    
    # Título
    pdf.set_font("Helvetica", 'B', 18)
    pdf.set_text_color(*COLOR_PRIMARIO)
    pdf.cell(0, 10, f"REPORTE GERENCIAL - LOTE #{lote.id_lote_}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.set_font("Helvetica", 'I', 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, f"Fecha de emision: {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(10)
    
    # --- Información del Lote ---
    pdf.set_font("Helvetica", 'B', 14)
    pdf.set_text_color(*COLOR_PRIMARIO)
    pdf.cell(0, 8, "1. Informacion del Lote", new_x="LMARGIN", new_y="NEXT")
    pdf.set_fill_color(*COLOR_GRIS)
    
    pdf.set_font("Helvetica", '', 11)
    pdf.set_text_color(0, 0, 0)
    
    # Simulamos columnas con celdas
    pdf.cell(45, 8, "Tipo de Ave:", border=1, fill=True)
    pdf.cell(50, 8, f"{lote.tipo_ave}", border=1)
    pdf.cell(45, 8, "Edad (dias):", border=1, fill=True)
    pdf.cell(50, 8, f"{lote.edad_dias}", border=1, new_x="LMARGIN", new_y="NEXT")
    
    pdf.cell(45, 8, "Proveedor:", border=1, fill=True)
    pdf.cell(50, 8, f"{lote.proveedor}", border=1)
    pdf.cell(45, 8, "Peso Inicial (g):", border=1, fill=True)
    pdf.cell(50, 8, f"{lote.peso_inicial}", border=1, new_x="LMARGIN", new_y="NEXT")
    
    pdf.cell(45, 8, "Aves Actuales:", border=1, fill=True)
    pdf.cell(50, 8, f"{lote.saldo_actual}", border=1)
    pdf.cell(45, 8, "Mortalidad Total:", border=1, fill=True)
    pdf.cell(50, 8, f"{lote.mortalidad_porcentaje}%", border=1, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)
    
    # --- Resumen Ambiental (Promedios) ---
    lecturas = LecturaSensor.query.filter_by(id_lote=lote.id_lote_).all()
    pdf.set_font("Helvetica", 'B', 14)
    pdf.set_text_color(*COLOR_PRIMARIO)
    pdf.cell(0, 8, "2. Resumen Ambiental (Promedios Historicos)", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    
    if lecturas:
        avg_temp = round(sum(l.temperatura for l in lecturas) / len(lecturas), 1)
        avg_hum = round(sum(l.humedad for l in lecturas) / len(lecturas), 1)
        avg_amon = round(sum(l.amoniaco for l in lecturas) / len(lecturas), 1)
        
        pdf.set_font("Helvetica", '', 11)
        pdf.cell(63, 8, f"Temperatura: {avg_temp} C", border=1, align='C')
        pdf.cell(63, 8, f"Humedad: {avg_hum}%", border=1, align='C')
        pdf.cell(64, 8, f"Amoniaco: {avg_amon} ppm", border=1, align='C', new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.set_font("Helvetica", 'I', 11)
        pdf.cell(0, 8, "No hay datos de telemetria suficientes.", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    # --- Registro de Vacunación ---
    pdf.set_font("Helvetica", 'B', 14)
    pdf.set_text_color(*COLOR_PRIMARIO)
    pdf.cell(0, 8, "3. Historial de Vacunacion", new_x="LMARGIN", new_y="NEXT")
    
    vacunas = Vacunacion.query.filter_by(id_lote=lote.id_lote_).all()
    if vacunas:
        pdf.set_font("Helvetica", 'B', 10)
        pdf.set_fill_color(*COLOR_SECUNDARIO)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(30, 8, "Fecha", border=1, fill=True)
        pdf.cell(60, 8, "Vacuna", border=1, fill=True)
        pdf.cell(60, 8, "Objetivo", border=1, fill=True)
        pdf.cell(40, 8, "Estado", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_font("Helvetica", '', 10)
        pdf.set_text_color(0, 0, 0)
        for i, v in enumerate(vacunas):
            fill = (i % 2 == 0)
            if fill: pdf.set_fill_color(*COLOR_GRIS)
            fecha_str = v.fecha_aplicacion.strftime('%Y-%m-%d') if v.fecha_aplicacion else 'N/A'
            nombre = v.vacuna.nombre if v.vacuna else 'Desconocida'
            obj = v.vacuna.enfermedad_objetivo if v.vacuna else 'Desconocido'
            pdf.cell(30, 8, fecha_str, border=1, fill=fill)
            pdf.cell(60, 8, nombre[:28], border=1, fill=fill)
            pdf.cell(60, 8, obj[:28], border=1, fill=fill)
            pdf.cell(40, 8, v.estado, border=1, fill=fill, new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.set_font("Helvetica", 'I', 11)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, "No se han registrado vacunas para este lote.", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)
    
    # --- Registro de Mortalidad ---
    pdf.set_font("Helvetica", 'B', 14)
    pdf.set_text_color(*COLOR_PRIMARIO)
    pdf.cell(0, 8, "4. Detalle de Mortalidad", new_x="LMARGIN", new_y="NEXT")
    
    bajas = Mortalidad.query.filter_by(id_lote_=lote.id_lote_).all()
    if bajas:
        pdf.set_font("Helvetica", 'B', 10)
        pdf.set_fill_color(*COLOR_SECUNDARIO)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(40, 8, "Fecha", border=1, fill=True)
        pdf.cell(100, 8, "Causa", border=1, fill=True)
        pdf.cell(50, 8, "Cantidad de Aves", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_font("Helvetica", '', 10)
        pdf.set_text_color(0, 0, 0)
        for i, b in enumerate(bajas):
            fill = (i % 2 == 0)
            if fill: pdf.set_fill_color(*COLOR_GRIS)
            pdf.cell(40, 8, str(b.fecha.strftime('%Y-%m-%d')), border=1, fill=fill)
            pdf.cell(100, 8, str(b.causa)[:45], border=1, fill=fill)
            pdf.cell(50, 8, str(b.cantidad), border=1, fill=fill, new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.set_font("Helvetica", 'I', 11)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, "No se han registrado bajas en este lote.", new_x="LMARGIN", new_y="NEXT")
        
    temp_dir = tempfile.gettempdir()
    base_filename = f'reporte_{lote.id_lote_}_{int(datetime.now().timestamp())}.pdf'
    pdf_path = os.path.join(temp_dir, base_filename)
    
    pdf.output(pdf_path)
    
    return send_file(pdf_path, as_attachment=True)
